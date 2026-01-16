import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import json
import time

# ==============================
# 🔑 Qwen API 配置
# ==============================
QWEN_API_KEY = "sk-a9a56d27bc1845349cc65ca500743d4c"
QWEN_MODEL = "qwen-max"
API_URL = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"

HEADERS = {
    "Authorization": f"Bearer {QWEN_API_KEY}",
    "Content-Type": "application/json"
}

def call_qwen(prompt: str, max_retries=3) -> str:
    """
    调用 DashScope Qwen API（正确格式）
    支持 qwen-max / qwen-plus / qwen-turbo
    """
    payload = {
        "model": QWEN_MODEL,
        "input": {
            "messages": [
                {"role": "user", "content": prompt}
            ]
        },
        "parameters": {
            "max_tokens": 500,
            "temperature": 0.3,
            "top_p": 0.8
        }
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(API_URL, headers=HEADERS, data=json.dumps(payload))
            if response.status_code == 200:
                result = response.json()
                # ✅ DashScope Qwen 返回的是 output.text，不是 choices！
                return result['output']['text'].strip()
            else:
                err_msg = response.json().get("message", response.text)
                print(f"❌ API 错误 (尝试 {attempt + 1}/{max_retries}): {response.status_code} - {err_msg}")
        except Exception as e:
            print(f"⚠️ 请求异常 (尝试 {attempt + 1}/{max_retries}): {e}")
        time.sleep(1)
    return "【答案】此处应由大模型生成，但 API 调用失败，请检查配置。"


def generate_scene_summary(scene_name: str, steps: list) -> str:
    steps_text = "\n".join([f"{i+1}. {step}" for i, step in enumerate(steps)])
    prompt = f"""你是一个 HR 系统专家。请根据以下测试步骤，总结出该场景的整体业务流程。

要求：
- 输出格式为一段连贯自然语言
- 不要编号、不要 bullet points
- 包含主要角色（如HR、用人部门）、关键环节、业务目的
- 长度控制在 300 字以内
- 不要编造未提及的功能

场景名称：{scene_name}
测试步骤列表：
{steps_text}

请输出总结："""
    return call_qwen(prompt)


# ================== 配置 ==================
EXCEL_FILE = "薪酬管理.xlsx"
OUTPUT_FOLDER = Path(EXCEL_FILE).stem
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ================== 创建 DOCX ==================
doc = Document()

# === 全局字体 ===
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# === 主标题 ===
main_title = Path(EXCEL_FILE).stem
title = doc.add_heading(main_title, level=0)
title_run = title.runs[0]
title_run.font.size = Pt(30)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ================== 处理每个 Sheet ==================
with pd.ExcelFile(EXCEL_FILE) as xls:
    for sheet_name in xls.sheet_names:
        print(f"🔄 处理工作表: {sheet_name}")

        df_raw = pd.read_excel(xls, sheet_name=sheet_name, header=None, dtype=str)
        df_raw = df_raw.fillna("")

        if df_raw.empty:
            print(f"⚠️ 跳过空表: {sheet_name}")
            continue

        nrows, ncols = df_raw.shape
        col_map = {}
        data_start_row = -1

        # === 找表头 ===
        for i in range(nrows):
            row = df_raw.iloc[i]
            non_empty_cols = [(j, str(row[j]).strip()) for j in range(ncols) if str(row[j]).strip() != ""]
            if not non_empty_cols:
                continue
            values = [v for _, v in non_empty_cols]
            indices = [j for j, _ in non_empty_cols]
            #哈哈哈改了
            target_fields = ["测试步骤", "功能路径", "输入数据/特殊信息"]#, "预期结果"
            if all(f in values for f in target_fields):
                for f in target_fields:
                    col_map[f] = indices[values.index(f)]
                data_start_row = i + 1
                break

        if not col_map or data_start_row == -1:
            print(f"⚠️ 未找到表头，跳过: {sheet_name}")
            continue

        # === 提取有效数据 ===
        data_rows = []
        for i in range(data_start_row, nrows):
            row = df_raw.iloc[i]
            # try:
            #     step = str(row[col_map["测试步骤"]]).strip()
            #     path = str(row[col_map["功能路径"]]).strip()
            #     input_data = str(row[col_map["输入数据/特殊信息"]]).strip()
            #     expected = str(row[col_map["预期结果"]]).strip()
            # except Exception:
            #     break
            # if not step or not path or not input_data or not expected:
            #     break
            # data_rows.append([step, path, input_data, expected])
            #哈哈哈改了
            try:
                step = str(row[col_map["测试步骤"]]).strip()
                path = str(row[col_map["功能路径"]]).strip()
                input_data = str(row[col_map["输入数据/特殊信息"]]).strip()
            except Exception:
                break
            if not step or not path or not input_data:
                break
            data_rows.append([step, path, input_data])  # 不再存 expected

        if not data_rows:
            print(f"⚠️ 无有效数据，跳过: {sheet_name}")
            continue

        # === 添加 Sheet 二级标题 ===
        sec_title = doc.add_heading(sheet_name, level=1)
        sec_title_run = sec_title.runs[0]
        sec_title_run.font.size = Pt(24)
        sec_title_run.font.bold = True
        sec_title_run.font.color.rgb = RGBColor(0, 0, 0)

        # === 1. 生成并插入【场景级问答对】（每个 Sheet 仅一次）===
        steps_list = [row[0] for row in data_rows]
        summary_answer = generate_scene_summary(sheet_name, steps_list)

        # 🔗 手动拼接测试步骤链（加在 AI 答案前面）
        steps_chain = " → ".join(steps_list)
        prefixed_answer = f"测试步骤如下：{steps_chain}。\n\n{summary_answer}"

        doc.add_paragraph(f"【问题】{sheet_name}的整体测试流程是什么？")
        doc.add_paragraph(f"【答案】{prefixed_answer}")  # ← 使用拼接后的答案
        doc.add_paragraph("───────────────────────────────────────")
        #哈哈哈改了
        #for idx, (step, path, input_data, expected) in enumerate(data_rows, 1):
        for idx, (step, path, input_data) in enumerate(data_rows, 1):
            # 【问题】行这里666可以加表头的

            #666改了
            doc.add_paragraph(f"【问题】如何测试{sheet_name}中的{step}？")

            answer_lines = [
                f"测试步骤：{step}",
                f"功能路径：{path.replace('->', ' → ')}",
                f"输入数据 / 特殊信息：{input_data if input_data else '-'}",
                #哈哈哈改了
                #f"预期结果：{expected}"
            ]
            full_answer = "\n".join(answer_lines)
            doc.add_paragraph("【答案】")
            doc.add_paragraph(full_answer)

            doc.add_paragraph("")

        doc.add_page_break()

# ================== 保存 ==================
output_file = os.path.join(OUTPUT_FOLDER, f"{main_title}.docx")
doc.save(output_file)
print(f"\n✅ 已生成合并文档: {output_file}")
print(f"🎉 输出目录: {OUTPUT_FOLDER}")