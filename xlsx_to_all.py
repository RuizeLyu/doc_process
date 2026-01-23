import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import json
import time
import csv

# ==============================
# 🔑 Qwen API 配置
# ==============================
QWEN_API_KEY = "sk-a9a56d27bc1845349cc65ca500743d4c"
QWEN_MODEL = "qwen-max"
API_URL = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"

HEADERS = {
    "Authorization": f"Bearer {QWEN_API_KEY}",
    "Content-Type": "application/json"
}

def call_qwen(prompt: str, max_retries=3) -> str:
    """
    调用 DashScope Qwen API（正确格式）
    支持 qwen-max / qwen-plus / qwen-turbo
    """
    payload = {
        "model": QWEN_MODEL,
        "input": {
            "messages": [
                {"role": "user", "content": prompt}
            ]
        },
        "parameters": {
            "max_tokens": 500,
            "temperature": 0.3,
            "top_p": 0.8
        }
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(API_URL, headers=HEADERS, data=json.dumps(payload))
            if response.status_code == 200:
                result = response.json()
                # ✅ DashScope Qwen 返回的是 output.text，不是 choices！
                return result['output']['text'].strip()
            else:
                err_msg = response.json().get("message", response.text)
                print(f"❌ API 错误 (尝试 {attempt + 1}/{max_retries}): {response.status_code} - {err_msg}")
        except Exception as e:
            print(f"⚠️ 请求异常 (尝试 {attempt + 1}/{max_retries}): {e}")
        time.sleep(1)
    return "【答案】此处应由大模型生成，但 API 调用失败，请检查配置。"

def generate_scene_summary(scene_name: str, steps: list) -> str:
    steps_text = "\n".join([f"{i+1}. {step}" for i, step in enumerate(steps)])
    prompt = f"""你是一个 HR 系统专家。请根据以下测试步骤，总结出该场景的整体业务流程。

要求：
- 输出格式为一段连贯自然语言
- 不要编号、不要 bullet points
- 包含主要角色（如HR、用人部门）、关键环节、业务目的
- 长度控制在 300 字以内
- 不要编造未提及的功能

场景名称：{scene_name}
测试步骤列表：
{steps_text}

请输出总结："""
    return call_qwen(prompt)

def process_excel_file(excel_file: str):
    """
    处理 Excel 文件，同时生成 DOCX 和 CSV 文件
    """
    print(f"\n🚀 开始处理 Excel 文件: {excel_file}")
    
    # 配置
    EXCEL_FILE = excel_file
    OUTPUT_FOLDER = "result"
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    
    main_title = Path(EXCEL_FILE).stem
    
    # 准备存储问答对的列表（用于 CSV）
    qa_pairs = []
    
    # 创建 DOCX
    doc = Document()
    
    # === 全局字体 ===
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # === 主标题 ===
    title = doc.add_heading(main_title, level=0)
    title_run = title.runs[0]
    title_run.font.size = Pt(30)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 处理每个 Sheet
    with pd.ExcelFile(EXCEL_FILE) as xls:
        for sheet_name in xls.sheet_names:
            print(f"🔄 处理工作表: {sheet_name}")

            df_raw = pd.read_excel(xls, sheet_name=sheet_name, header=None, dtype=str)
            df_raw = df_raw.fillna("")

            if df_raw.empty:
                print(f"⚠️ 跳过空表: {sheet_name}")
                continue

            nrows, ncols = df_raw.shape
            col_map = {}
            data_start_row = -1

            # === 找表头 ===
            for i in range(nrows):
                row = df_raw.iloc[i]
                non_empty_cols = [(j, str(row[j]).strip()) for j in range(ncols) if str(row[j]).strip() != ""]
                if not non_empty_cols:
                    continue
                values = [v for _, v in non_empty_cols]
                indices = [j for j, _ in non_empty_cols]
                
                # 尝试不同的表头格式
                target_fields_list = [
                    ["测试角色", "测试步骤", "功能路径", "输入数据/特殊信息", "预期结果"],
                    ["测试角色", "测试步骤", "功能路径", "输入数据", "预期结果"],
                    ["测试步骤", "功能路径", "输入数据/特殊信息", "预期结果"],
                    ["测试步骤", "功能路径", "输入数据", "预期结果"],
                    ["测试步骤", "功能路径", "输入数据/特殊信息"],
                    ["测试步骤", "功能路径", "输入数据"],
                    ["关键环节", "操作角色", "操作步骤", "输入数据", "预期结果"],
                    ["测试角色", "测试步骤", "输入数据", "预期结果"],
                    ["操作角色", "测试步骤", "功能路径", "输入数据", "预期结果"]
                ]
                
                found = False
                for target_fields in target_fields_list:
                    if all(f in values for f in target_fields):
                        for f in target_fields:
                            col_map[f] = indices[values.index(f)]
                        data_start_row = i + 1
                        found = True
                        break
                
                if found:
                    break

            if not col_map or data_start_row == -1:
                print(f"⚠️ 未找到表头，跳过: {sheet_name}")
                continue

            # === 提取有效数据 ===
            data_rows = []
            for i in range(data_start_row, nrows):
                row = df_raw.iloc[i]
                try:
                    # 尝试获取测试角色，可能有不同的字段名
                    if "测试角色" in col_map:
                        test_role = str(row[col_map["测试角色"]]).strip()
                    elif "操作角色" in col_map:
                        test_role = str(row[col_map["操作角色"]]).strip()
                    else:
                        test_role = ""
                    
                    # 尝试获取测试步骤，可能有不同的字段名
                    if "测试步骤" in col_map:
                        step = str(row[col_map["测试步骤"]]).strip()
                    elif "操作步骤" in col_map:
                        step = str(row[col_map["操作步骤"]]).strip()
                    elif "关键环节" in col_map:
                        step = str(row[col_map["关键环节"]]).strip()
                    else:
                        step = ""
                    
                    # 尝试获取功能路径，可能不存在
                    if "功能路径" in col_map:
                        path = str(row[col_map["功能路径"]]).strip()
                    else:
                        path = ""
                    
                    # 尝试获取输入数据，可能有不同的字段名
                    if "输入数据/特殊信息" in col_map:
                        input_data = str(row[col_map["输入数据/特殊信息"]]).strip()
                    elif "输入数据" in col_map:
                        input_data = str(row[col_map["输入数据"]]).strip()
                    else:
                        input_data = ""
                    
                    # 尝试获取预期结果，可能不存在
                    if "预期结果" in col_map:
                        expected = str(row[col_map["预期结果"]]).strip()
                    else:
                        expected = ""

                except Exception:
                    break
                if not step:
                    break
                data_rows.append([test_role, step, path, input_data, expected])  # 包含所有字段，包括测试角色

            if not data_rows:
                print(f"⚠️ 无有效数据，跳过: {sheet_name}")
                continue

            # === 添加 Sheet 二级标题到 DOCX ===
            sec_title = doc.add_heading(sheet_name, level=1)
            sec_title_run = sec_title.runs[0]
            sec_title_run.font.size = Pt(24)
            sec_title_run.font.bold = True
            sec_title_run.font.color.rgb = RGBColor(0, 0, 0)

            # === 1. 生成并插入【场景级问答对】（每个 Sheet 仅一次）===
            steps_list = [row[1] for row in data_rows]  # row[1] 是测试步骤
            summary_answer = generate_scene_summary(sheet_name, steps_list)

            # 🔗 手动拼接测试步骤链（加在 AI 答案前面）
            steps_chain = " → ".join(steps_list)
            prefixed_answer = f"测试步骤如下：{steps_chain}。\n\n{summary_answer}"

            # 添加到 DOCX
            doc.add_paragraph(f"【问题】{sheet_name}的整体测试流程是什么？")
            doc.add_paragraph(f"【答案】{prefixed_answer}")

            # 添加到 CSV
            qa_pairs.append({
                "问题": f"{sheet_name}的整体测试流程是什么？",
                "预期回答": prefixed_answer
            })

            # === 2. 生成并插入【步骤级问答对】（每个步骤一条）===
            for idx, (test_role, step, path, input_data, expected) in enumerate(data_rows, 1):
                # 构建问题
                question = f"如何测试{sheet_name}中的{step}？"
                
                # 构建答案
                answer_lines = []
                # 如果有测试角色，添加到答案中
                if test_role:
                    answer_lines.append(f"测试角色：{test_role}")
                answer_lines.extend([
                    f"测试步骤：{step}",
                    f"功能路径：{path.replace('->', ' → ')}",
                    f"输入数据 / 特殊信息：{input_data if input_data else '-'}",
                    f"预期结果：{expected if expected else '-'}"
                ])
                full_answer = "\n".join(answer_lines)
                
                # 添加到 DOCX
                doc.add_paragraph(f"【问题】{question}")
                doc.add_paragraph("【答案】")
                doc.add_paragraph(full_answer)
                doc.add_paragraph("")
                
                # 添加到 CSV
                qa_pairs.append({
                    "问题": question,
                    "预期回答": full_answer
                })

            doc.add_page_break()

    # ================== 保存 DOCX ==================
    docx_file = os.path.join(OUTPUT_FOLDER, f"{main_title}.docx")
    
    # 🔥 关键修复：如果文件已存在，先删除，避免尾部垃圾字节
    if os.path.exists(docx_file):
        os.remove(docx_file)

    doc.save(docx_file)
    print(f"\n✅ 已生成 DOCX 文件: {docx_file}")

    # ================== 保存 CSV ==================
    csv_file = os.path.join(OUTPUT_FOLDER, f"{main_title}_qa.csv")
    print(f"📄 准备写入 CSV 文件: {csv_file}")

    # 写入 CSV 文件，使用 UTF-8-SIG 编码以支持中文
    try:
        # 确保目录存在
        os.makedirs(os.path.dirname(csv_file), exist_ok=True)
        
        # 尝试删除旧文件（如果存在）
        if os.path.exists(csv_file):
            os.remove(csv_file)
            print(f"🔄 已删除旧文件")
        
        with open(csv_file, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=["问题", "预期回答"])
            writer.writeheader()
            writer.writerows(qa_pairs)
        
        print(f"✅ 已生成 CSV 文件: {csv_file}")
        print(f"🎉 共收集 {len(qa_pairs)} 条问答对")
    except Exception as e:
        print(f"\n❌ 生成 CSV 文件失败: {e}")
        # 尝试使用当前目录作为备选
        alt_csv_file = os.path.join(os.getcwd(), f"{main_title}_qa.csv")
        print(f"🔄 尝试使用备选路径: {alt_csv_file}")
        try:
            with open(alt_csv_file, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.DictWriter(f, fieldnames=["问题", "预期回答"])
                writer.writeheader()
                writer.writerows(qa_pairs)
            print(f"✅ 已生成 CSV 文件 (备选路径): {alt_csv_file}")
            print(f"🎉 共收集 {len(qa_pairs)} 条问答对")
        except Exception as e2:
            print(f"❌ 备选路径写入失败: {e2}")

    print(f"\n🎉 处理完成: {excel_file}")

def main():
    """
    主函数
    """
    # 处理所有 Excel 文件
    excel_files = ["files/人力助手.xls"]
    
    for excel_file in excel_files:
        if os.path.exists(excel_file):
            process_excel_file(excel_file)
        else:
            print(f"❌ 找不到文件: {excel_file}")

if __name__ == "__main__":
    main()
